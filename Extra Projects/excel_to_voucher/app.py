import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
import datetime
import win32print
import win32api
from num2words import num2words
import time
import os
import re
import threading

def replace_text_in_doc(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
    return doc

def generate_voucher(template_bytes, employee_data, index):
    doc = Document(template_bytes)
    voucher_no = f"VOUCHER-{int(index) + 1:03d}"
    current_date = datetime.datetime.now().strftime("%Y-%m-%d")
    
    try:
        salary_val = float(employee_data.get("Total", 0))
        salary_words = num2words(salary_val, lang='en_IN').replace("-", " ")
    except Exception:
        salary_words = ""
    
    replacements = {
        "NO.": voucher_no,
        "DATE :": f"DATE : {current_date}",
        "PAYEE : cash": f"PAYEE : {employee_data.get('Payee', 'cash')}",
        "Bombaybhel": employee_data.get("StationName", ""),
        "Afzal .": employee_data.get("EmployeeName", ""),
        "Helper": employee_data.get("Designation", ""),
        "29,161": str(employee_data.get("Total", "")),
        "1006015": employee_data.get("EmployeeCode", ""),
        "twenty nine thousand one hundred and sixty one": salary_words,
    }
    
    doc = replace_text_in_doc(doc, replacements)
    
    voucher_file = BytesIO()
    doc.save(voucher_file)
    voucher_file.seek(0)
    return voucher_file, doc

def print_document(file_path):
    printer_name = win32print.GetDefaultPrinter()
    win32api.ShellExecute(0, "print", file_path, f'/d:"{printer_name}"', ".", 0)

def threaded_print(voucher_files, save_dir):
    def print_job(file_path):
        print_document(file_path)
    
    threads = []
    for name, file_obj in voucher_files:
        temp_path = os.path.join(save_dir, f"{name}_voucher.docx")
        with open(temp_path, "wb") as f:
            f.write(file_obj.getvalue())
        t = threading.Thread(target=print_job, args=(temp_path,))
        threads.append(t)
        t.start()
    
    for t in threads:
        t.join()

def main():
    st.title("Voucher Generator App")
    excel_file = st.file_uploader("Upload Employee Data (Excel)", type=["xlsx", "xls"])
    voucher_template = st.file_uploader("Upload Voucher Template (Word .docx)", type=["docx"])
    
    if excel_file and voucher_template:
        df = pd.read_excel(excel_file)
        column_mapping = {col.strip().lower(): col for col in df.columns}
        name_column = column_mapping.get("employeename", None)

        if not name_column:
            st.error("Column 'EmployeeName' not found in Excel. Please check column names.")
            return
        
        df.index = df.index.astype(str)
        search_suggestions = df[name_column].astype(str).tolist()
        search_query = st.selectbox("Search Employee by Name", options=[""] + search_suggestions, index=0)
        if search_query:
            search_query = search_query.strip().lower()
            df = df[df[name_column].astype(str).str.lower().str.contains(search_query, na=False) |
                    df['Designation'].astype(str).str.lower().str.contains(search_query, na=False)]
        
        all_selected = st.checkbox("Select All Employees")
        selected_rows = st.multiselect("Select Employees", df.index.tolist(), format_func=lambda x: str(df.loc[x, name_column]) if x in df.index else "Unknown", default=df.index.tolist() if all_selected else [])
        
        cancel_printing = st.checkbox("Cancel Printing")
        save_dir = st.text_input("Enter Save Directory", value="generated_vouchers")
        os.makedirs(save_dir, exist_ok=True)
        
        if selected_rows:
            template_bytes = voucher_template.read()
            voucher_files = []
            progress_bar = st.progress(0)
            
            for i, index in enumerate(selected_rows):
                if cancel_printing:
                    st.warning("Printing Cancelled!")
                    return
                
                employee_data = df.loc[index].to_dict()
                safe_name = f"{employee_data.get('EmployeeName', f'employee_{index}')}_{employee_data.get('Designation', 'Unknown')}"
                safe_name = re.sub(r'[<>:"/\\|?*]', '_', safe_name).replace(' ', '_')
                voucher_file, doc = generate_voucher(BytesIO(template_bytes), employee_data, index)
                voucher_files.append((safe_name, voucher_file))
                progress_bar.progress((i + 1) / len(selected_rows))
                
                if i < 3 and st.checkbox(f"Preview {safe_name}"):
                    st.download_button(f"Download {safe_name} Voucher", data=voucher_file, file_name=f"{safe_name}_voucher.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button("Print Only"):
                    threaded_print(voucher_files, save_dir)
                    st.success("All Vouchers Sent to Printer!")
            
            with col2:
                if st.button("Save Only"):
                    zip_buffer = BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        for name, file_obj in voucher_files:
                            zip_file.writestr(f"{name}_voucher.docx", file_obj.getvalue())
                    zip_buffer.seek(0)
                    st.download_button("Download ZIP", data=zip_buffer, file_name="vouchers.zip", mime="application/zip")
            
            with col3:
                if st.button("Print and Save"):
                    threaded_print(voucher_files, save_dir)
                    st.success("Vouchers Printed and Saved Successfully!")
            
            with col4:
                if st.button("Cancel"):
                    st.warning("Operation Cancelled")
                    return

if __name__ == "__main__":
    main()
